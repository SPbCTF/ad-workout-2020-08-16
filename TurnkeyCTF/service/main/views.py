from flask import Flask, request, session, redirect, url_for, render_template, flash, jsonify, send_file
from . forms import  SignUpForm, SignInForm
from main import app
from difflib import SequenceMatcher
from flask_pymongo import PyMongo 
from bson import ObjectId
from stop_words import get_stop_words
from flask_session import Session as FSession
from docx import Document
import sys
import random
import string
import imaplib
import hashlib
import time
import os
import subprocess
import re
import logging
import json
import datetime
import urllib.parse
import bcrypt


app.config['pepper'] = "AEju5_P6B9mEVuhcxk&vufpuz?Lk#UH?9v*^XaS*!Tm8%DHY79e3L=mxT4tKMkpKq__DkaGhYuFCkUhES6e#Jzk!+&uzVwBt&v$"   #<<-- change this
app.config['SECRET_KEY'] = 'jkdsgfh87w45tu834jeqT@#$5234#$2#R567j56967u45679084j76$^&%3456'                                    #<<-- change this

app.config['DEBUG'] = True
app.config['threaded']=True
app.config['MONGO_URI'] = "mongodb://mongo_db:27017/turnkey"
app.config['SESSION_PERMANENT'] = True
app.config['SESSION_TYPE']= "mongodb"
logging.basicConfig(level=logging.INFO)
mongo = PyMongo(app)
pepper = app.config['pepper']
app.config['SESSION_MONGODB'] = mongo.db.flask_session
FSession(app)


@app.after_request
def add_headers(response):
    response.headers["X-Frame-Options"] = "SAMEORIGIN"
    response.headers["X-XSS-Protection"]="1; mode=block"
    response.headers["X-Content-Type-Options"]="nosniff"
    response.headers["X-Frame-Options"]="DENY"
    #response.headers["Cache-control"]="no-store"
    #response.headers["Pragma"]="no-cache"
    response.headers["Server"]="no server for you, dear hacker"
    return response


@app.route('/') # main page. redirect to tasks viewer is user available 
def index():
    if user_available():
        return redirect(url_for('task_viewer'))
    else:
        flash('You are not authenticated','error')
        session['user_available']=False
    return redirect(url_for('signin'))

def task_template(): 
    tasks_raw = mongo.db.tasks.find({ 'Project': { '$exists': False}}, {"Name":1, "Category":1, "_id":1})[:100]
    tasks_list=[]
    def sortlast(val): 
        return val[2]
    for task in tasks_raw:
        tasks_list.append([str(task["_id"]), task["Name"], task["Category"]])
    tasks_list.sort(key = sortlast, reverse=True)
    return([tasks_list])


def user_available():
    try:
        if session['user_available']:
            return True
        else:
            return False
    except:
        return False 


@app.route('/CTF_templator', methods=['GET']) #CTF template page
def CTF_templator():
    if user_available():
        task_list = task_template()
        projects = mongo.db.users.find_one({'username': session['current_user']}, {'Projects.Name':1, "_id":0})
        users = mongo.db.users.find({}, {'username':1, "_id":0})
        if projects:
            return render_template('CTF_templator.html', tasks=task_list[0], projects=map(lambda d: d['Name'], projects.get("Projects")), users=map(lambda d: d['username'], users))
        return render_template('CTF_templator.html', tasks=task_list[0], projects=['Create at least one project'], users=map(lambda d: d['username'], users))
    flash('You are not authenticated','error')
    return redirect(url_for('signin'))


@app.route('/task_viewer', methods=['GET']) #tasks view page
def task_viewer():
    if user_available():
        task_list = task_template() 
        return render_template('task_viewer.html', tasks=task_list[0]) 
    flash('You are not authenticated','error')
    return redirect(url_for('signin'))


@app.route('/edit_task', methods=['GET']) #task edit page
def edit_task():
    if user_available():
        task_list = task_template() 
        return render_template('edit_task.html', tasks=task_list[0])
    flash('You are not authenticated','error')
    return redirect(url_for('signin'))


@app.route('/add_task', methods=['GET','POST']) # task adding page
def add_task():
    if user_available():
        return render_template('add_task.html')
    flash('You are not authenticated','error')
    return redirect(url_for('signin'))


@app.route('/signup', methods=['GET', 'POST'])
def signup():
    if user_available():
        return redirect(url_for('task_viewer'))
    signupform = SignUpForm(request.form)
    if request.method == 'POST' and signupform.validate_on_submit():
        if mongo.db.users.find_one({'username': signupform.username.data})==None:
            salt = bcrypt.gensalt()
            password = bcrypt.hashpw((pepper+signupform.password.data).encode('utf-8'), salt)
            result = mongo.db.users.insert_one({'firstname': signupform.firstname.data, 'lastname': signupform.lastname.data, 'username': signupform.username.data, 'password': password, 'email': signupform.email.data, 'salt': salt})
            return redirect(url_for('signin'))
        else:
            flash("Something wrong",'error')
            return render_template('signup.html', signupform=signupform)
    return render_template('signup.html', signupform=signupform)


@app.route('/signin', methods=['GET', 'POST'])
def signin():
    if user_available():
        return redirect(url_for('task_viewer'))
    signinform = SignInForm(request.form)
    if request.method == 'POST':
        if signinform.validate_on_submit():
            username = signinform.username.data
            log = mongo.db.users.find_one({'username': signinform.username.data})
            if log!=None:
                if log["password"] == bcrypt.hashpw((pepper+signinform.password.data).encode('utf-8'), log["salt"]):
                    current_user = log["username"]
                    session['current_user'] = current_user
                    session['lang']="ru"
                    session['user_available'] = True
                    return redirect(url_for('task_viewer'))
                else:
                    flash('Wrong login or password','error') 
            else:
                flash('Wrong login or password','error')
    return render_template('signin.html', signinform=signinform)


@app.route('/logout')
def logout():
    if user_available():
        session['user_available'] = False
    return redirect(url_for('signin'))


@app.route('/get_info', methods=['GET'])
def get_info():
    task_id = request.args.get('id')
    try:
        ObjectId(task_id)
    except:
        return "Error"
    task = mongo.db.tasks.find_one({'_id': ObjectId(task_id)})
    task_json={"Category":task["Category"],"Category_full":task["Category_full"],'Name':task["Name"],'Description': task["Description"],"Flag":task["Flag"]}
    return jsonify(result=task_json)


@app.route('/add_task_handler', methods=['POST']) #add task to database
def add_task_handler():
    if user_available():
        Category_list = {"Web", "Rev", "HTB", "Steg", "Cryp", "Misc" , "Joy", "PWN"}
        Category_list_full =  {"Web", "Reverse", "HTB", "Steganography", "Cryptography", "Misc" , "Joy", "PWN"}
        if request.form['Name'] and request.form['Description'] and request.form['Flag'] and request.form['Category'] and request.form['Category_full']:
            if (request.form['Category'] in Category_list) and (request.form['Category_full'] in Category_list_full):
                mongo.db.tasks.insert_one({'Name': request.form['Name'], 'Description': request.form['Description'], 'Flag': request.form['Flag'], 'Category': request.form['Category'], 'Category_full': request.form['Category_full']})
                return jsonify(result={"message":'Success',"category":'success'})
            return jsonify(result={"message":'Wrong category',"category":'error'})
        return jsonify(result={"message":'Fields should not be empty',"category":'error'})
    flash('You are not authenticated','error')
    return redirect(url_for('signin'))


@app.route('/edit_task_handler', methods=['POST']) # edit task
def edit_task_handler():
    if user_available():
        try:
            Category_list = {"Web", "Rev", "HTB", "Steg", "Cryp", "Misc" , "Joy", "PWN"}
            Category_list_full =  {"Web", "Reverse", "HTB", "Steganography", "Cryptography", "Misc" , "Joy", "PWN"}
            if request.form['id'] and request.form['Name'] and request.form['Description'] and request.form['Flag'] and request.form['Category'] and request.form['Category_full']:
                if (request.form['Category'] in Category_list) and (request.form['Category_full'] in Category_list_full):
                    mongo.db.tasks.update_one({'_id': ObjectId(request.form['id'])}, {"$set":{'Name': request.form['Name'], 'Description': request.form['Description'], 'Flag': request.form['Flag'], 'Category': request.form['Category'], 'Category_full': request.form['Category_full']}})
                    return jsonify(result={"message":"Successfully edited!","category":'success'})
                return jsonify(result={"message":'Bad category',"category":'error'})
            return jsonify(result={"message":'Fields should not be empty',"category":'error'})
        except:
            return jsonify(result={"message":'Something wrong',"category":'error'})
    flash('You are not authenticated','error')
    return redirect(url_for('signin'))


@app.route('/delete_task_handler', methods=['POST']) #delete task
def delete_task_handler():
    if user_available():
        if request.form['id']:
            mongo.db.tasks.remove( { "_id":ObjectId(request.form['id'])} )
            return jsonify(result={"message":"Successfully deleted!","category":'success'})
        return jsonify(result={"message":"Wrong ID!","category":'error'})
    flash('You are not authenticated','error')
    return redirect(url_for('signin'))


@app.route('/search_handler', methods=['GET'])  # search
def search_handler():

    def sortSecond(val): 
        return val[1]  #used for extracting the second value from each list of lists

    if user_available():
        Search_string = request.args.get('Search_string')
        search_options = request.args.get('Search_options')
        if search_options==False:
            return jsonify(result={"message":"Select at least 1 option!","category":'error'})

        Search_string = re.sub(r'[^a-zа-я0-9 ]', ' ', Search_string.lower())
        search_word_split = Search_string.split()   
        stop_words = get_stop_words('english')+get_stop_words('russian') 
        search_word_split = [w for w in search_word_split if not w in stop_words] 
        search_options = search_options.split(',')
        options_map = ["Name","Description", "Flag"]
        query_args = {}
        for option in search_options:
            if option in options_map:
                query_args[option] = 1
            else:
                return jsonify(result={"message":"Wrong options!","category":'error'})
        tasks_raw = mongo.db.tasks.find({}, query_args)
        sorted_list=[]      
        for task in tasks_raw:
            match=0.0
            task_search = re.sub(r'[^a-zа-я0-9 ]', ' ', ' '.join(task[option] for option in search_options).lower())
            tasks_word_split = task_search.lower().split()
            tasks_word_split = [w for w in tasks_word_split if not w in stop_words]
            for search_word in search_word_split:
                match_word = 0.0
                for task_word in tasks_word_split:
                    match_temp = SequenceMatcher(lambda x: x == " ",search_word, task_word).ratio()
                    if match_temp>match:
                        match_word=match_temp
                match+=match_word
            match=match/len(search_word_split)
            if match!=0.0:
                sorted_list.append((str(task["_id"]),match))
        end = time.time()
        sorted_list.sort(key = sortSecond)
        sorted_list = [item[0] for item in sorted_list]
        return jsonify(result=sorted_list[::-1])

    return jsonify(result={"message":"Something wrong","category":'error'})
    flash('You are not authenticated','error')
    return redirect(url_for('signin'))


@app.route('/save_task', methods=['POST']) #add task to database
def save_task():
    if user_available():
        Category_list = {"Web", "Rev", "HTB", "Steg", "Cryp", "Misc" , "Joy", "PWN"}
        Category_list_full =  {"Web", "Reverse", "HTB", "Steganography", "Cryptography", "Misc" , "Joy", "PWN"}
        if request.form['id'] and request.form['Name'] and request.form['Description'] and request.form['Flag'] and request.form['Category'] and request.form['Category_full'] and request.form['Project']:
            if (request.form['Category'] in Category_list) and (request.form['Category_full'] in Category_list_full):
                projects = mongo.db.users.find_one({'username': session['current_user']}, {'Projects':1, "_id":0})
                if projects !=None:
                    project = mongo.db.users.find_one({'username': session['current_user'], "Projects.Name": request.form['Project']},{'Projects' : { '$elemMatch': {"Name":request.form['Project']}},'Projects._id':1, "_id":0})
                    if project!=None:
                        if mongo.db.projects.find_one({'_id': ObjectId(project.get("Projects")[0].get("_id")),"tasks._id":ObjectId(request.form['id'])}) == None:
                            new_task = mongo.db.tasks.insert_one({'Name': request.form['Name'], 'Description': request.form['Description'], 'Flag': request.form['Flag'], 'Category': request.form['Category'], 'Category_full': request.form['Category_full'], 'Project':ObjectId(project.get("Projects")[0].get("_id"))})
                            mongo.db.projects.update_one({'_id': ObjectId(project.get("Projects")[0].get("_id"))}, {'$push': {'tasks': {'_id' : new_task.inserted_id}}})  
                        else:
                            mongo.db.tasks.update({'_id': ObjectId(request.form['id'])}, {'$set': {'Name': request.form['Name'], 'Description': request.form['Description'], 'Flag': request.form['Flag'], 'Category': request.form['Category'], 'Category_full': request.form['Category_full'], 'Project': project.get("Projects")[0].get("_id")}})
                        return jsonify(result={"message":'Success',"category":'success'})
                    project_id = mongo.db.projects.insert_one({'Name': request.form['Project']})
                    mongo.db.users.update_one({'username': session['current_user']},{'$push': {"Projects":{"Name":request.form['Project'],'_id':ObjectId(project_id.inserted_id)}}})
                    new_task = mongo.db.tasks.insert_one({'Name': request.form['Name'], 'Description': request.form['Description'], 'Flag': request.form['Flag'], 'Category': request.form['Category'], 'Category_full': request.form['Category_full'], 'Project':project_id.inserted_id})
                    mongo.db.projects.update_one({'_id': project_id.inserted_id}, {'$push': {'tasks': {'_id' : new_task.inserted_id}}})
                    return jsonify(result={"message":'Success',"category":'success'})
                project_id = mongo.db.projects.insert_one({'Name': request.form['Project'], "tasks":[{'id':request.form['id'], 'Name': request.form['Name'], 'Description': request.form['Description'], 'Flag': request.form['Flag'], 'Category': request.form['Category'], 'Category_full': request.form['Category_full']}]})
                mongo.db.users.update_one({'username': session['current_user']},{'$set': {"Projects":[{"Name":request.form['Project'],'id':ObjectId(project_id.inserted_id)}]}})
                return jsonify(result={"message":'Success',"category":'success'})
            return jsonify(result={"message":'Bad category',"category":'error'})
        return jsonify(result={"message":'Fields should not be empty',"category":'error'})
    flash('You are not authenticated','error')
    return redirect(url_for('signin'))


@app.route('/project_selection', methods=['GET'])
def project_selection():
    if user_available():
        if 'project' in request.args:
            projects = mongo.db.users.find_one({'username': session['current_user']}, {'Projects':1, "_id":0})
            if projects !=None:
                project = mongo.db.users.find_one({'username': session['current_user'], "Projects.Name": request.args.get('project')},{'Projects' : { '$elemMatch': {"Name":request.args.get('project')}},'Projects._id':1, "_id":0})
                if project!=None:
                    tasks = mongo.db.projects.find_one({'_id': ObjectId(project.get("Projects")[0].get("_id"))}, {'tasks':1, "_id":0})
                    tasks = list(mongo.db.tasks.find({'_id':{ '$in': list(map(lambda d: d['_id'], tasks.get("tasks")))}},{'Name':1, "_id":1,"Category":1}))
                    for task in tasks:
                        task['_id'] = str(task['_id'])
                    return jsonify(result=tasks)
            return jsonify(result={"message":'New project created',"category":'success'})
        return jsonify(result={"message":'Fields should not be empty',"category":'error'})
    flash('You are not authenticated','error')
    return redirect(url_for('signin'))


@app.route('/create_CTF', methods=['GET'])
def create_CTF():
    if user_available():
        project = mongo.db.users.find_one({'username': session['current_user'], "Projects.Name": request.args.get('project')}, {'Projects' : { '$elemMatch': {"Name":request.args.get('project')}},'Projects._id':1, "_id":0})
        if project ==None:
           print("F")
        else:
            document = Document()
            fname = "/app/main/projects/"+request.args.get('project')+".docx"
            tasks = mongo.db.projects.find_one({'_id': ObjectId(project.get("Projects")[0].get("_id"))}, {'tasks':1, "_id":0})
            tasks = mongo.db.tasks.find({'_id':{ '$in': list(map(lambda d: d['_id'], tasks.get("tasks")))}},{'Name':1, "_id":0,"Category":1,"Description":1,"Flag":1})
            tasks = list(tasks)
            document.add_heading('Tasks for '+ request.args.get('project')+' CTF')
            for task in tasks:
                document.add_paragraph('Taskname:')
                document.add_paragraph(task['Name'])
                document.add_paragraph('Category:')
                document.add_paragraph(task['Category'])
                document.add_paragraph('Description:')
                document.add_paragraph(task['Description'])
                document.add_paragraph('Flag:')
                document.add_paragraph(task['Flag'])            
            document.save(fname)
    return send_file(fname, attachment_filename=request.args.get('project')+".docx")


@app.route('/user_selection', methods=['GET'])
def user_selection():
    if user_available():
        if 'user' in request.args and 'project' in request.args:
            projects = mongo.db.users.find_one({'username': session['current_user']}, {'Projects':1, "_id":0})
            if projects !=None:
                project = mongo.db.users.find_one({'username': session['current_user'], "Projects.Name": request.args.get('project')},{'Projects' : { '$elemMatch': {"Name":request.args.get('project')}},'Projects._id':1, "_id":0})
                if project!=None:
                    user = mongo.db.users.find_one({'username': request.args.get('user')})
                    if user!=None:
                        mongo.db.users.update_one({'username': request.args.get('user')},{'$push': {"Projects":{"Name":request.args.get('project'),'_id':ObjectId(project.get("Projects")[0].get("_id"))}}})
                        return jsonify(result={"message":'Success',"category":'success'})
                    return jsonify(result={"message":'No such user',"category":'error'})
                return jsonify(result={"message":'You dont have such project',"category":'error'})
            return jsonify(result={"message":'You dont have any projects',"category":'error'})
        return jsonify(result={"message":'Fields should not be empty',"category":'error'})


if __name__ == '__main__':
    app.run()
